# TESTE EXPORTERS FINAL

from __future__ import annotations

from io import BytesIO
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

from config import FUNDO_PADRAO


def dataframe_to_excel_bytes(df: pd.DataFrame, sheet_name: str = "dados") -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
        workbook = writer.book
        worksheet = writer.sheets[sheet_name]

        wrap = workbook.add_format({
            "text_wrap": True,
            "valign": "top",
            "border": 1,
        })
        header = workbook.add_format({
            "bold": True,
            "text_wrap": True,
            "valign": "vcenter",
            "align": "center",
            "border": 1,
            "bg_color": "#D9E2F3",
        })

        for idx, col in enumerate(df.columns):
            if df.empty:
                max_len = 10
            else:
                max_len = int(df[col].astype(str).str.len().max())
            width = min(max(len(str(col)), max_len) + 2, 45)
            worksheet.set_column(idx, idx, width, wrap)

        for col_idx, col in enumerate(df.columns):
            worksheet.write(0, col_idx, col, header)

        worksheet.freeze_panes(1, 0)
        worksheet.autofilter(0, 0, max(len(df), 1), max(len(df.columns) - 1, 0))

    return output.getvalue()


def build_box_covers_docx_bytes(records: list[dict]) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    for idx, row in enumerate(records):
        if idx > 0:
            doc.add_section(WD_SECTION.NEW_PAGE)

        p = doc.add_paragraph()
        p.style = doc.styles["Title"]
        run = p.add_run("CAPA / ETIQUETA DE CAIXA")
        run.bold = True
        run.font.size = Pt(18)

        info = [
            ("Fundo", row.get("fundo") or FUNDO_PADRAO),
            ("Grupo", row.get("grupo", "")),
            ("Subgrupo", row.get("subgrupo", "")),
            ("Série", row.get("serie", "")),
            ("Subsérie", row.get("subserie", "")),
            ("Dossiê / Processo", row.get("dossie_processo", "")),
            ("Item documental", row.get("item_documental", "")),
            ("Setor", row.get("setor", "")),
            ("Caixa", row.get("caixa", "")),
            ("Datas-limite", row.get("datas_limite", "")),
            (
                "Temporalidade",
                f"Corrente: {row.get('prazo_corrente', '')} | "
                f"Intermediário: {row.get('prazo_intermediario', '')} | "
                f"Destinação: {row.get('destinacao_final', '')}",
            ),
        ]

        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in info:
            cells = table.add_row().cells
            cells[0].text = str(label)
            cells[1].text = str(value or "-")

        doc.add_paragraph(
            "Universidade do Estado de Santa Catarina - UDESC / Centro de Ciências Tecnológicas - CCT"
        )

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _doc_template(buffer: BytesIO) -> SimpleDocTemplate:
    return SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )


def _styles():
    styles = getSampleStyleSheet()

    title = ParagraphStyle(
        "title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        alignment=1,
    )
    subtitle = ParagraphStyle(
        "subtitle",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        alignment=1,
    )
    body = ParagraphStyle(
        "body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
    )
    body_center = ParagraphStyle(
        "body_center",
        parent=body,
        alignment=1,
    )
    body_bold = ParagraphStyle(
        "body_bold",
        parent=body,
        fontName="Helvetica-Bold",
    )
    small = ParagraphStyle(
        "small",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=7,
        leading=9,
    )
    small_bold = ParagraphStyle(
        "small_bold",
        parent=small,
        fontName="Helvetica-Bold",
    )
    return {
        "title": title,
        "subtitle": subtitle,
        "body": body,
        "body_center": body_center,
        "body_bold": body_bold,
        "small": small,
        "small_bold": small_bold,
    }


def _safe_text(text: str) -> str:
    return (
        str(text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(_safe_text(text or "-"), style)


def _field_box(label: str, value: str, label_style: ParagraphStyle, value_style: ParagraphStyle):
    return [
        Paragraph(_safe_text(label), label_style),
        Paragraph(_safe_text(value or "-"), value_style),
    ]


def _compute_codigo(row: dict) -> str:
    codigo = row.get("codigo_classificacao") or row.get("item_documental_codigo") or ""
    if codigo:
        return str(codigo)

    parts = [
        row.get("grupo"),
        row.get("subgrupo"),
        row.get("serie"),
        row.get("subserie"),
    ]
    parts = [str(v) for v in parts if v]
    return " / ".join(parts)


def _compute_especificacao(row: dict) -> str:
    quantidade = row.get("quantidade", "")
    caixa = row.get("caixa", "")
    espec = f"{quantidade} caixa(s)" if quantidade else "caixa-arquivo"
    if caixa:
        espec += f" | Caixa {caixa}"
    return espec


def _compute_total_mensuracao(records: list[dict]) -> float:
    total = 0.0
    for row in records:
        qtd = row.get("quantidade") or 0
        try:
            total += float(qtd) * 0.14
        except Exception:
            pass
    return total


def _compute_date_limits(records: list[dict]) -> str:
    values = []
    for row in records:
        val = str(row.get("datas_limite", "")).strip()
        if val:
            values.append(val)
    if not values:
        return "-"
    return "; ".join(values)


def _compute_document_summary(records: list[dict]) -> str:
    nomes = []
    for row in records:
        nome = str(row.get("tipo_documental", "")).strip()
        if nome and nome not in nomes:
            nomes.append(nome)
    if not nomes:
        return "-"
    return "; ".join(nomes)


def build_elimination_listing_dataframe(records: list[dict]) -> pd.DataFrame:
    rows = []
    for idx, row in enumerate(records, start=1):
        quantidade = row.get("quantidade", "")
        rows.append({
            "Nº": f"{idx:02d}",
            "Código de Classificação": _compute_codigo(row),
            "Nome do Documento": row.get("tipo_documental", ""),
            "Datas-limite": row.get("datas_limite", ""),
            "Unidade de Arquivamento - Quantidade": quantidade,
            "Unidade de Arquivamento - Especificação": _compute_especificacao(row),
            "Observações e/ou Justificativas": row.get("justificativa") or row.get("observacoes", ""),
        })
    return pd.DataFrame(rows)


def build_elimination_pdf(records: list[dict], meta: dict | None = None) -> bytes:
    meta = meta or {}
    buffer = BytesIO()
    doc = _doc_template(buffer)
    s = _styles()

    story = [
        Paragraph("LISTAGEM DE ELIMINAÇÃO DE DOCUMENTOS", s["title"]),
        Spacer(1, 0.25 * cm),
    ]

    header_table = Table(
        [
            [
                _field_box("ÓRGÃO/ENTIDADE", meta.get("orgao_entidade", ""), s["body_bold"], s["body"]),
                _field_box(
                    "LISTAGEM Nº/ANO",
                    f"{meta.get('listagem_numero', '')}/{meta.get('listagem_ano', '')}",
                    s["body_bold"],
                    s["body"],
                ),
            ],
            [
                _field_box("UNIDADE/SETOR", meta.get("unidade_setor", ""), s["body_bold"], s["body"]),
                _field_box("", "", s["body_bold"], s["body"]),
            ],
        ],
        colWidths=[13.8 * cm, 4.0 * cm],
    )
    header_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.extend([header_table, Spacer(1, 0.18 * cm)])

    data = [[
        Paragraph("CÓDIGO DE CLASSIFICAÇÃO", s["small_bold"]),
        Paragraph("NOME DO DOCUMENTO", s["small_bold"]),
        Paragraph("DATAS-LIMITE", s["small_bold"]),
        Paragraph("UNIDADE DE ARQUIVAMENTO\nQuantidade", s["small_bold"]),
        Paragraph("UNIDADE DE ARQUIVAMENTO\nEspecificação", s["small_bold"]),
        Paragraph("OBSERVAÇÃO", s["small_bold"]),
    ]]

    total_mensuracao = _compute_total_mensuracao(records)

    for row in records:
        qtd = row.get("quantidade") or ""
        data.append([
            Paragraph(_safe_text(_compute_codigo(row) or "-"), s["body"]),
            Paragraph(_safe_text(str(row.get("tipo_documental", "") or "-")), s["body"]),
            Paragraph(_safe_text(str(row.get("datas_limite", "") or "-")), s["body"]),
            Paragraph(_safe_text(str(qtd or "-")), s["body_center"]),
            Paragraph(_safe_text(_compute_especificacao(row) or "-"), s["body"]),
            Paragraph(_safe_text(str(row.get("justificativa") or row.get("observacoes") or "-")), s["body"]),
        ])

    for _ in range(max(0, 10 - len(records))):
        data.append([
            Paragraph("", s["body"]),
            Paragraph("", s["body"]),
            Paragraph("", s["body"]),
            Paragraph("", s["body"]),
            Paragraph("", s["body"]),
            Paragraph("", s["body"]),
        ])

    listing_table = Table(
        data,
        repeatRows=1,
        colWidths=[3.0 * cm, 5.0 * cm, 2.2 * cm, 1.6 * cm, 2.8 * cm, 3.8 * cm],
    )
    listing_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#efefef")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (3, 1), (3, -1), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.extend([listing_table, Spacer(1, 0.18 * cm)])

    mens_table = Table(
        [[
            Paragraph("Mensuração total:", s["body_bold"]),
            Paragraph(f"{total_mensuracao:.2f} metros lineares", s["body"]),
        ]],
        colWidths=[4.5 * cm, 13.3 * cm],
    )
    mens_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.extend([mens_table, Spacer(1, 0.18 * cm)])

    story.append(
        Paragraph(
            "O quadro abaixo somente deverá ser preenchido se os documentos a serem eliminados necessitarem de comprovação de aprovação das contas pelos órgãos competentes.",
            s["small"],
        )
    )

    contas_table = Table(
        [
            [
                Paragraph("Conta(s) do(s) exercício(s) de:", s["small_bold"]),
                Paragraph("Conta(s) aprovada(s) pelo órgão competente em:", s["small_bold"]),
                Paragraph("Documento oficial que registra a aprovação, órgão que aprovou, data e meio de divulgação:", s["small_bold"]),
            ],
            [Paragraph("", s["small"]), Paragraph("", s["small"]), Paragraph("", s["small"])],
            [Paragraph("", s["small"]), Paragraph("", s["small"]), Paragraph("", s["small"])],
        ],
        colWidths=[4.2 * cm, 4.8 * cm, 8.8 * cm],
    )
    contas_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.extend([contas_table, Spacer(1, 0.18 * cm)])

    data_local = meta.get("data_local") or datetime.now().strftime("%d/%m/%Y")
    local_data_table = Table(
        [[
            Paragraph("LOCAL E DATA", s["body_bold"]),
            Paragraph(f"{meta.get('local', '')}, {data_local}", s["body"]),
        ]],
        colWidths=[4.0 * cm, 13.8 * cm],
    )
    local_data_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.extend([local_data_table, Spacer(1, 0.18 * cm)])

    sign_table = Table(
        [[
            [
                Paragraph("RESPONSÁVEL PELO ÓRGÃO", s["body_bold"]),
                Paragraph(str(meta.get("responsavel_orgao", "") or ""), s["body"]),
                Paragraph(str(meta.get("cargo_responsavel_orgao", "") or ""), s["body"]),
            ],
            [
                Paragraph("PRESIDENTE DA CPAD", s["body_bold"]),
                Paragraph(str(meta.get("presidente_cpad", "") or ""), s["body"]),
                Paragraph(str(meta.get("cargo_presidente_cpad", "") or ""), s["body"]),
            ],
            [
                Paragraph("RESPONSÁVEL PELA SELEÇÃO", s["body_bold"]),
                Paragraph(str(meta.get("responsavel_selecao", "") or ""), s["body"]),
                Paragraph(str(meta.get("cargo_responsavel_selecao", "") or ""), s["body"]),
            ],
        ]],
        colWidths=[5.93 * cm, 5.93 * cm, 5.93 * cm],
    )
    sign_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 16),
    ]))
    story.append(sign_table)

    doc.build(story)
    return buffer.getvalue()


def build_edital_pdf(records: list[dict], meta: dict | None = None) -> bytes:
    meta = meta or {}
    buffer = BytesIO()
    doc = _doc_template(buffer)
    s = _styles()

    listagem_ref = f"{meta.get('listagem_numero', '')}/{meta.get('listagem_ano', '')}".strip("/")
    edital_ref = f"{meta.get('edital_numero', '')}/{meta.get('edital_ano', '')}".strip("/")
    datas_limite = meta.get("datas_limite_gerais") or _compute_date_limits(records)
    conjuntos = meta.get("conjuntos_documentais") or _compute_document_summary(records)

    titulo = f"EDITAL DE CIÊNCIA DE ELIMINAÇÃO DE DOCUMENTOS Nº {edital_ref}"
    texto = (
        f"O {meta.get('cargo_titular_orgao', '')} e o Presidente da Comissão Permanente de Avaliação - CPAD, "
        f"conforme Processo {meta.get('processo_numero', '')} e considerando a Listagem de Eliminação de "
        f"Documentos nº {listagem_ref}, fazem saber a quem possa interessar, que a contar do período de 30 "
        f"(trinta) dias corridos, subsequente à data de publicação deste Edital, se não houver contestações, "
        f"o {meta.get('orgao_entidade', '')} eliminará os documentos relativos a {conjuntos}, do período "
        f"{datas_limite}, do {meta.get('unidade_setor', '')}. Os interessados, no prazo citado, poderão "
        f"requerer, às suas expensas e mediante petição dirigida à CPAD do(a) {meta.get('orgao_entidade', '')}, "
        f"a retirada ou cópias de documentos, avulsos ou processos, bem como o desentranhamento ou cópias "
        f"de folhas de um processo."
    )

    story = [
        Paragraph("ANEXO II", s["subtitle"]),
        Spacer(1, 0.15 * cm),
        Paragraph("MODELO DE EDITAL DE CIÊNCIA DE ELIMINAÇÃO DE DOCUMENTOS", s["subtitle"]),
        Spacer(1, 0.35 * cm),
        Paragraph(titulo, s["title"]),
        Spacer(1, 0.35 * cm),
        Paragraph(_safe_text(texto), s["body"]),
        Spacer(1, 1.2 * cm),
    ]

    sign_table = Table(
        [[
            [
                Paragraph(str(meta.get("nome_titular_orgao", "") or ""), s["body_center"]),
                Paragraph(str(meta.get("cargo_titular_orgao", "") or ""), s["body_center"]),
            ],
            [
                Paragraph(str(meta.get("presidente_cpad", "") or ""), s["body_center"]),
                Paragraph("Presidente da CPAD", s["body_center"]),
            ],
        ]],
        colWidths=[8.9 * cm, 8.9 * cm],
    )
    sign_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 12),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
    ]))
    story.append(sign_table)

    doc.build(story)
    return buffer.getvalue()


def build_termo_pdf(records: list[dict], meta: dict | None = None) -> bytes:
    meta = meta or {}
    buffer = BytesIO()
    doc = _doc_template(buffer)
    s = _styles()

    listagem_ref = f"{meta.get('listagem_numero', '')}/{meta.get('listagem_ano', '')}".strip("/")
    edital_ref = f"{meta.get('edital_numero', '')}/{meta.get('edital_ano', '')}".strip("/")
    total_mensuracao = _compute_total_mensuracao(records)
    datas_limite = meta.get("datas_limite_gerais") or _compute_date_limits(records)
    data_eliminacao_extenso = meta.get("data_eliminacao_extenso", "")
    if not data_eliminacao_extenso:
        data_eliminacao_extenso = datetime.now().strftime("%d/%m/%Y")

    texto = (
        f"Aos {data_eliminacao_extenso}, o(a) {meta.get('orgao_entidade', '')}, de acordo com o que consta "
        f"na Listagem de Eliminação de Documentos {listagem_ref} e respectivo Edital de Ciência de Eliminação "
        f"de Documentos {edital_ref}, publicado no Diário Oficial do Estado {meta.get('doe_numero_data', '')}, "
        f"conforme processo {meta.get('processo_numero', '')}, procedeu à eliminação de "
        f"{total_mensuracao:.2f} metros lineares de documentos integrantes do acervo documental do(a) "
        f"{meta.get('unidade_setor', '')}, do período relativo a {datas_limite}."
    )

    story = [
        Paragraph("ANEXO III", s["subtitle"]),
        Spacer(1, 0.15 * cm),
        Paragraph("MODELO DE TERMO DE ELIMINAÇÃO DE DOCUMENTOS", s["subtitle"]),
        Spacer(1, 0.35 * cm),
        Paragraph("TERMO DE ELIMINAÇÃO DE DOCUMENTOS", s["title"]),
        Spacer(1, 0.35 * cm),
        Paragraph(_safe_text(texto), s["body"]),
        Spacer(1, 1.2 * cm),
    ]

    sign_table = Table(
        [[
            [
                Paragraph(str(meta.get("nome_titular_orgao", "") or ""), s["body_center"]),
                Paragraph(str(meta.get("cargo_titular_orgao", "") or ""), s["body_center"]),
            ],
            [
                Paragraph(str(meta.get("presidente_cpad", "") or ""), s["body_center"]),
                Paragraph("Presidente da CPAD", s["body_center"]),
            ],
        ]],
        colWidths=[8.9 * cm, 8.9 * cm],
    )
    sign_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 12),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
    ]))
    story.append(sign_table)

    doc.build(story)
    return buffer.getvalue()