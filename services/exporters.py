from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from io import BytesIO
from pathlib import Path

import pandas as pd
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
)

from config import FUNDO_PADRAO


def dataframe_to_excel_bytes(df: pd.DataFrame, sheet_name: str = "dados") -> bytes:
    """
    Exporta DataFrame para Excel em memória.
    Corrigido para colunas vazias / NaN.
    """
    output = BytesIO()

    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)

        workbook = writer.book
        worksheet = writer.sheets[sheet_name]

        fmt_header = workbook.add_format({
            "bold": True,
            "text_wrap": True,
            "valign": "vcenter",
            "align": "center",
            "border": 1,
            "bg_color": "#D9E2F3",
        })

        fmt_body = workbook.add_format({
            "text_wrap": True,
            "valign": "top",
            "border": 1,
        })

        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, fmt_header)

        for idx, col in enumerate(df.columns):
            if df.empty:
                max_len = 10
            else:
                serie = df[col].fillna("").astype(str)
                max_len_val = serie.str.len().max()
                max_len = int(max_len_val) if pd.notna(max_len_val) else 10

            largura = min(max(len(str(col)), max_len) + 2, 45)
            worksheet.set_column(idx, idx, largura, fmt_body)

        worksheet.freeze_panes(1, 0)

        if len(df.columns) > 0:
            worksheet.autofilter(
                0,
                0,
                max(len(df), 1),
                len(df.columns) - 1
            )

    output.seek(0)
    return output.getvalue()


def _replace_text_in_paragraph(paragraph, mapping: dict) -> None:
    """
    Substitui placeholders no formato {{ campo }} em um parágrafo.
    Mantém ao máximo a formatação do primeiro run.
    """
    if not paragraph.text:
        return

    texto = paragraph.text
    alterado = False

    for key, value in mapping.items():
        placeholder = "{{ " + key + " }}"
        if placeholder in texto:
            texto = texto.replace(placeholder, str(value or ""))
            alterado = True

    if alterado:
        if paragraph.runs:
            paragraph.runs[0].text = texto
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = texto


def _replace_text_in_table(table, mapping: dict) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_text_in_paragraph(paragraph, mapping)
            for nested_table in cell.tables:
                _replace_text_in_table(nested_table, mapping)


def _replace_placeholders_in_doc(doc: Document, mapping: dict) -> None:
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, mapping)

    for table in doc.tables:
        _replace_text_in_table(table, mapping)


def build_box_covers_from_template_docx_bytes(records: list[dict]) -> bytes:
    """
    Gera capas/etiquetas de caixa a partir de um template DOCX institucional.
    O layout é preservado a partir do arquivo:
    data/reference/etiqueta_caixas.docx
    """
    base_dir = Path(__file__).resolve().parent.parent
    template_path = base_dir / "data" / "reference" / "etiqueta_caixas.docx"

    if not template_path.exists():
        raise FileNotFoundError(f"Template DOCX não encontrado: {template_path}")

    output_doc = Document()

    if output_doc.paragraphs:
        p = output_doc.paragraphs[0]._element
        p.getparent().remove(p)

    for index, row in enumerate(records):
        temp_doc = Document(str(template_path))

        mapping = {
            "fundo": row.get("fundo") or FUNDO_PADRAO,
            "unidade_macro": row.get("unidade_macro", ""),
            "setor_responsavel": row.get("setor_responsavel", ""),
            "unidade_documentacao": row.get("unidade_documentacao", ""),
            "codigo_classificacao": row.get("codigo_classificacao", ""),
            "assunto": row.get("assunto", ""),
            "datas_limite_resumo": row.get("datas_limite_resumo", ""),
            "datas_limite_detalhadas": row.get("datas_limite_detalhadas", ""),
            "prazo_corrente": row.get("prazo_corrente", ""),
            "prazo_intermediario": row.get("prazo_intermediario", ""),
            "destinacao": row.get("destinacao", ""),
            "destaque_permanente": row.get("destaque_permanente", ""),
            "numero_caixa": row.get("numero_caixa", ""),
            "observacao": row.get("observacao", ""),
            "lista_itens_caixa": row.get("lista_itens_caixa", ""),
        }

        _replace_placeholders_in_doc(temp_doc, mapping)

        for element in temp_doc.element.body:
            output_doc.element.body.append(deepcopy(element))

        if index < len(records) - 1:
            output_doc.add_page_break()

    output = BytesIO()
    output_doc.save(output)
    output.seek(0)
    return output.getvalue()


def build_box_covers_docx_bytes(records: list[dict]) -> bytes:
    """
    Mantida por compatibilidade com o restante do projeto.
    Agora usa o template DOCX institucional.
    """
    return build_box_covers_from_template_docx_bytes(records)


def _doc_template(buffer: BytesIO) -> SimpleDocTemplate:
    return SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )


def build_elimination_listing_dataframe(records: list[dict]) -> pd.DataFrame:
    rows = []
    for idx, row in enumerate(records, start=1):
        codigo = row.get("codigo_classificacao") or row.get("item_documental_codigo") or ""
        if not codigo:
            codigo = " / ".join(
                v for v in [
                    row.get("grupo"),
                    row.get("subgrupo"),
                    row.get("serie"),
                    row.get("subserie"),
                ] if v
            )

        prazo = (
            f"Corrente: {row.get('prazo_corrente', '')} | "
            f"Intermediário: {row.get('prazo_intermediario', '')}"
        )

        quantidade = row.get("quantidade", "")
        caixa = row.get("caixa", "")
        espec = f"{quantidade} caixa(s)" if quantidade else "caixa-arquivo"
        if caixa:
            espec += f" | Caixa {caixa}"

        rows.append({
            "Nº": f"{idx:02d}",
            "Código de Classificação": codigo,
            "Nome do Documento": row.get("tipo_documental", ""),
            "Datas-limite": row.get("datas_limite", ""),
            "Unidade de Arquivamento - Quantidade": quantidade,
            "Unidade de Arquivamento - Especificação": espec,
            "Prazo de Guarda": prazo,
            "Destinação": row.get("destinacao_final", ""),
            "Observações e/ou Justificativas": row.get("justificativa") or row.get("observacoes", ""),
        })

    return pd.DataFrame(rows)


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    safe = (
        str(text or "-")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )
    return Paragraph(safe, style)


def build_elimination_pdf(records: list[dict], meta: dict | None = None) -> bytes:
    meta = meta or {}
    buffer = BytesIO()
    doc = _doc_template(buffer)

    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        alignment=1,
    )
    body = ParagraphStyle(
        "body",
        parent=styles["BodyText"],
        fontSize=8,
        leading=10,
    )
    small = ParagraphStyle(
        "small",
        parent=styles["BodyText"],
        fontSize=7,
        leading=9,
    )

    story = [
        Paragraph("LISTAGEM DE ELIMINAÇÃO DE DOCUMENTOS", title),
        Spacer(1, 0.25 * cm),
    ]

    header_table = Table(
        [
            [
                _p(
                    f"ÓRGÃO/ENTIDADE: {meta.get('orgao_entidade', '')}",
                    body,
                ),
                _p(
                    f"LISTAGEM Nº/ANO: {meta.get('listagem_numero', '')}/{meta.get('listagem_ano', '')}",
                    body,
                ),
            ],
            [
                _p(
                    f"UNIDADE/SETOR: {meta.get('unidade_setor', '')}",
                    body,
                ),
                _p("", body),
            ],
        ],
        colWidths=[13.8 * cm, 4.0 * cm],
    )
    header_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.extend([header_table, Spacer(1, 0.18 * cm)])

    data = [[
        _p("CÓDIGO DE CLASSIFICAÇÃO", small),
        _p("NOME DO DOCUMENTO", small),
        _p("DATAS-LIMITE", small),
        _p("UNIDADE DE ARQUIVAMENTO - Qtd", small),
        _p("UNIDADE DE ARQUIVAMENTO Especificação", small),
        _p("OBSERVAÇÃO", small),
    ]]

    total_mensuracao = 0.0
    for row in records:
        qtd = row.get("quantidade") or 0
        try:
            total_mensuracao += float(qtd) * 0.14
        except Exception:
            pass

        espec = "caixa-arquivo"
        if row.get("caixa"):
            espec += f" | Caixa {row.get('caixa')}"

        obs = row.get("justificativa") or row.get("observacoes") or ""

        codigo = row.get("codigo_classificacao") or row.get("item_documental_codigo") or ""
        if not codigo:
            codigo = " / ".join(
                v for v in [
                    row.get("grupo"),
                    row.get("subgrupo"),
                    row.get("serie"),
                    row.get("subserie"),
                ] if v
            )

        data.append([
            _p(codigo, body),
            _p(row.get("tipo_documental", ""), body),
            _p(row.get("datas_limite", ""), body),
            _p(str(qtd or ""), body),
            _p(espec, body),
            _p(obs, body),
        ])

    for _ in range(max(0, 10 - len(records))):
        data.append([
            _p("", body),
            _p("", body),
            _p("", body),
            _p("", body),
            _p("", body),
            _p("", body),
        ])

    listing_table = Table(
        data,
        repeatRows=1,
        colWidths=[3.0 * cm, 5.0 * cm, 2.2 * cm, 1.4 * cm, 2.8 * cm, 4.0 * cm],
    )
    listing_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#efefef")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.extend([listing_table, Spacer(1, 0.18 * cm)])

    story.append(_p(f"Mensuração total: {total_mensuracao:.2f} metros lineares", body))
    story.append(Spacer(1, 0.18 * cm))
    story.append(_p(
        "(O quadro abaixo somente deverá ser preenchido se os documentos a serem eliminados "
        "necessitarem de comprovação de aprovação das contas pelos órgãos competentes.)",
        small,
    ))

    contas_table = Table(
        [
            [
                _p("Conta(s) do(s) exercício(s) de:", small),
                _p("Conta(s) aprovada(s) pelo órgão competente em:", small),
                _p(
                    "Documento Oficial que registra a aprovação, órgão que aprovou, data e meio de divulgação:",
                    small,
                ),
            ],
            [_p("", small), _p("", small), _p("", small)],
            [_p("", small), _p("", small), _p("", small)],
        ],
        colWidths=[4.2 * cm, 4.8 * cm, 8.8 * cm],
    )
    contas_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.extend([contas_table, Spacer(1, 0.18 * cm)])

    data_local = meta.get("data_local") or datetime.now().strftime("%d/%m/%Y")
    local_data_table = Table(
        [[_p(f"LOCAL E DATA: {meta.get('local', '')}, {data_local}", body)]],
        colWidths=[17.8 * cm],
    )
    local_data_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.extend([local_data_table, Spacer(1, 0.18 * cm)])

    sign_table = Table(
        [[
            _p(
                f"RESPONSÁVEL PELO ÓRGÃO{meta.get('responsavel_orgao', '')}{meta.get('cargo_responsavel_orgao', '')}",
                body,
            ),
            _p(
                f"PRESIDENTE DA CPAD: {meta.get('presidente_cpad', '')}{meta.get('cargo_presidente_cpad', '')}",
                body,
            ),
            _p(
                f"RESPONSÁVEL PELA SELEÇÃO{meta.get('responsavel_selecao', '')}{meta.get('cargo_responsavel_selecao', '')}",
                body,
            ),
        ]],
        colWidths=[5.93 * cm, 5.93 * cm, 5.93 * cm],
    )
    sign_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 16),
    ]))
    story.append(sign_table)

    doc.build(story)
    return buffer.getvalue()


def build_edital_pdf(records: list[dict], meta: dict | None = None) -> bytes:
    meta = meta or {}
    buffer = BytesIO()
    doc = _doc_template(buffer)

    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "title_edital",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        alignment=1,
    )
    body = ParagraphStyle(
        "body_edital",
        parent=styles["BodyText"],
        fontSize=10,
        leading=14,
        alignment=4,
    )
    body_left = ParagraphStyle(
        "body_edital_left",
        parent=styles["BodyText"],
        fontSize=10,
        leading=14,
    )

    edital_numero = f"{meta.get('edital_numero', '')}/{meta.get('edital_ano', '')}".strip("/")
    orgao = meta.get("orgao_entidade", "")
    unidade = meta.get("unidade_setor", "")
    titular = meta.get("nome_titular_orgao", "") or meta.get("responsavel_orgao", "")
    cargo_titular = meta.get("cargo_titular_orgao", "") or meta.get("cargo_responsavel_orgao", "")
    presidente_cpad = meta.get("presidente_cpad", "")
    conjuntos_documentais = meta.get("conjuntos_documentais", "")
    datas_limite = meta.get("datas_limite_gerais", "")
    processo_numero = meta.get("processo_numero", "")
    doe_numero_data = meta.get("doe_numero_data", "")
    local = meta.get("local", "")
    data_local = meta.get("data_local") or datetime.now().strftime("%d/%m/%Y")

    qtd_itens = len(records)

    story = [
        Paragraph("EDITAL DE CIÊNCIA DE ELIMINAÇÃO DE DOCUMENTOS", title),
        Spacer(1, 0.35 * cm),
    ]

    if edital_numero:
        story.append(Paragraph(f"EDITAL Nº {edital_numero}", body_left))
        story.append(Spacer(1, 0.2 * cm))

    texto = (
        f"O(A) {cargo_titular} {titular}, de acordo com a Listagem de Eliminação de Documentos "
        f"nº {meta.get('listagem_numero', '')}/{meta.get('listagem_ano', '')}, aprovada pelo(a) "
        f"{presidente_cpad}, faz saber a quem possa interessar que, se não houver oposição, o "
        f"{orgao} / {unidade} eliminará {qtd_itens} item(ns) documental(is)"
    )

    if conjuntos_documentais:
        texto += f", referentes a {conjuntos_documentais}"

    if datas_limite:
        texto += f", com datas-limite {datas_limite}"

    texto += (
        ". Os interessados, no prazo de 30 dias corridos a contar da data de publicação deste "
        "Edital, poderão requerer às suas expensas e mediante petição dirigida à Comissão "
        "Permanente de Avaliação de Documentos - CPAD, o desentranhamento ou cópias de documentos, "
        f"avulsos ou processos, bem como o desmembramento de folhas de um processo. Processo: {processo_numero}."
    )

    if doe_numero_data:
        texto += f" Publicação de referência: {doe_numero_data}."

    story.append(Paragraph(texto, body))
    story.append(Spacer(1, 1.0 * cm))

    story.append(Paragraph(f"{local}, {data_local}", body_left))
    story.append(Spacer(1, 1.2 * cm))

    if titular:
        story.append(Paragraph(f"{titular}", body_left))
    if cargo_titular:
        story.append(Paragraph(cargo_titular, body_left))

    doc.build(story)
    return buffer.getvalue()


def build_termo_pdf(records: list[dict], meta: dict | None = None) -> bytes:
    meta = meta or {}
    buffer = BytesIO()
    doc = _doc_template(buffer)

    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "title_termo",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        alignment=1,
    )
    body = ParagraphStyle(
        "body_termo",
        parent=styles["BodyText"],
        fontSize=10,
        leading=14,
        alignment=4,
    )
    body_left = ParagraphStyle(
        "body_termo_left",
        parent=styles["BodyText"],
        fontSize=10,
        leading=14,
    )

    edital_numero = f"{meta.get('edital_numero', '')}/{meta.get('edital_ano', '')}".strip("/")
    orgao = meta.get("orgao_entidade", "")
    unidade = meta.get("unidade_setor", "")
    responsavel = meta.get("nome_titular_orgao", "") or meta.get("responsavel_orgao", "")
    cargo_responsavel = meta.get("cargo_titular_orgao", "") or meta.get("cargo_responsavel_orgao", "")
    presidente_cpad = meta.get("presidente_cpad", "")
    data_eliminacao = meta.get("data_eliminacao_extenso", "") or meta.get("data_local", "")
    local = meta.get("local", "")
    processo_numero = meta.get("processo_numero", "")
    doe_numero_data = meta.get("doe_numero_data", "")

    qtd_itens = len(records)
    total_mensuracao = 0.0
    for row in records:
        qtd = row.get("quantidade") or 0
        try:
            total_mensuracao += float(qtd) * 0.14
        except Exception:
            pass

    story = [
        Paragraph("TERMO DE ELIMINAÇÃO DE DOCUMENTOS", title),
        Spacer(1, 0.35 * cm),
    ]

    texto = (
        f"Aos {data_eliminacao}, o {orgao} / {unidade}, de acordo com o que consta da "
        f"Listagem de Eliminação de Documentos nº {meta.get('listagem_numero', '')}/{meta.get('listagem_ano', '')}, "
        f"aprovada pelo(a) {presidente_cpad}, e do Edital de Ciência de Eliminação de Documentos "
        f"nº {edital_numero}, publicado"
    )

    if doe_numero_data:
        texto += f" em {doe_numero_data}"
    else:
        texto += " na forma regulamentar"

    texto += (
        ", procedeu à eliminação de documentos relativos a "
        f"{qtd_itens} item(ns) documental(is), totalizando aproximadamente "
        f"{total_mensuracao:.2f} metros lineares. Processo: {processo_numero}."
    )

    story.append(Paragraph(texto, body))
    story.append(Spacer(1, 1.0 * cm))

    story.append(Paragraph(f"{local}, {meta.get('data_local') or datetime.now().strftime('%d/%m/%Y')}", body_left))
    story.append(Spacer(1, 1.2 * cm))

    if responsavel:
        story.append(Paragraph(f"{responsavel}", body_left))
    if cargo_responsavel:
        story.append(Paragraph(cargo_responsavel, body_left))

    doc.build(story)
    return buffer.getvalue()